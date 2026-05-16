"""
Flask web application for the AI Text Humanizer.

Supports:
  - Paste plain text and get humanized output with before/after stats
  - Upload a .docx file with editable metadata controls
"""

import os
import uuid
import random
import tempfile
from datetime import datetime, timedelta

from flask import (
    Flask,
    render_template,
    request,
    jsonify,
    send_file,
)
from docx import Document

from humanizer import humanize_text, humanize_text_with_stats

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024

UPLOAD_DIR = os.path.join(tempfile.gettempdir(), "humanizer_uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)


def apply_metadata(doc, form_data):
    """
    Apply user-specified metadata to the document.
    Falls back to realistic random defaults if fields are empty.
    """
    props = doc.core_properties

    props.author = form_data.get("meta_author", "").strip()
    props.last_modified_by = form_data.get("meta_last_modified_by", "").strip() or props.author
    props.title = form_data.get("meta_title", "").strip()
    props.subject = form_data.get("meta_subject", "").strip()
    props.comments = ""
    props.category = ""
    props.keywords = ""
    props.content_status = ""
    props.identifier = ""
    props.language = form_data.get("meta_language", "en-US").strip() or "en-US"
    props.version = ""

    rev_str = form_data.get("meta_revision", "").strip()
    if rev_str:
        try:
            props.revision = max(1, int(rev_str))
        except ValueError:
            props.revision = random.randint(15, 80)
    else:
        props.revision = random.randint(15, 80)

    created_str = form_data.get("meta_created", "").strip()
    modified_str = form_data.get("meta_modified", "").strip()

    now = datetime.utcnow()

    if created_str:
        try:
            props.created = datetime.fromisoformat(created_str)
        except (ValueError, TypeError):
            props.created = now - timedelta(hours=random.randint(1, 72), minutes=random.randint(0, 59))
    else:
        props.created = now - timedelta(hours=random.randint(1, 72), minutes=random.randint(0, 59))

    if modified_str:
        try:
            props.modified = datetime.fromisoformat(modified_str)
        except (ValueError, TypeError):
            props.modified = now - timedelta(minutes=random.randint(1, 30))
    else:
        props.modified = now - timedelta(minutes=random.randint(1, 30))

    if props.modified and props.created and props.modified < props.created:
        props.modified = props.created + timedelta(hours=random.randint(1, 8))

    return doc


def strip_tracked_changes(doc):
    """Remove tracked changes / revision marks from the document."""
    try:
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        body = doc.element.body
        for tag in ['w:ins', 'w:del', 'w:rPrChange', 'w:pPrChange',
                    'w:sectPrChange', 'w:tblPrChange']:
            for elem in body.findall('.//' + tag, nsmap):
                parent = elem.getparent()
                if parent is not None:
                    if tag == 'w:ins':
                        for child in list(elem):
                            parent.insert(list(parent).index(elem), child)
                    parent.remove(elem)
    except Exception:
        pass
    return doc


def _fmt_dt(value):
    if not value:
        return ""
    try:
        return value.strftime("%Y-%m-%dT%H:%M")
    except AttributeError:
        return str(value)


def read_metadata(doc):
    """Read core document properties and return a JSON-friendly dict."""
    props = doc.core_properties
    return {
        "author": props.author or "",
        "last_modified_by": props.last_modified_by or "",
        "title": props.title or "",
        "subject": props.subject or "",
        "keywords": props.keywords or "",
        "category": props.category or "",
        "comments": props.comments or "",
        "content_status": props.content_status or "",
        "identifier": props.identifier or "",
        "language": props.language or "",
        "version": props.version or "",
        "revision": props.revision or 0,
        "created": _fmt_dt(props.created),
        "modified": _fmt_dt(props.modified),
        "last_printed": _fmt_dt(props.last_printed),
    }


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/inspect-docx", methods=["POST"])
def inspect_docx_api():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    if not file.filename or not file.filename.lower().endswith(".docx"):
        return jsonify({"error": "Please upload a .docx file"}), 400

    try:
        doc = Document(file.stream)
        return jsonify({"metadata": read_metadata(doc)})
    except Exception as e:
        return jsonify({"error": f"Failed to read document: {str(e)}"}), 500


@app.route("/api/humanize-text", methods=["POST"])
def humanize_text_api():
    data = request.get_json(silent=True)
    if not data or "text" not in data:
        return jsonify({"error": "No text provided"}), 400

    text = data["text"].strip()
    if not text:
        return jsonify({"error": "Text is empty"}), 400

    intensity = float(data.get("intensity", 0.5))
    result = humanize_text_with_stats(text, intensity=intensity)
    return jsonify(result)


@app.route("/api/humanize-docx", methods=["POST"])
def humanize_docx_api():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    if not file.filename or not file.filename.lower().endswith(".docx"):
        return jsonify({"error": "Please upload a .docx file"}), 400

    intensity = float(request.form.get("intensity", 0.5))

    uid = uuid.uuid4().hex
    input_path = os.path.join(UPLOAD_DIR, f"{uid}_input.docx")
    output_path = os.path.join(UPLOAD_DIR, f"{uid}_output.docx")
    file.save(input_path)

    try:
        doc = Document(input_path)

        for para in doc.paragraphs:
            if para.text.strip():
                humanized = humanize_text(para.text, intensity=intensity)
                if para.runs:
                    para.runs[0].text = humanized
                    for run in para.runs[1:]:
                        run.text = ""
                else:
                    para.text = humanized

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            humanized = humanize_text(para.text, intensity=intensity)
                            if para.runs:
                                para.runs[0].text = humanized
                                for run in para.runs[1:]:
                                    run.text = ""
                            else:
                                para.text = humanized

        doc = apply_metadata(doc, request.form)
        doc = strip_tracked_changes(doc)

        doc.save(output_path)
    except Exception as e:
        return jsonify({"error": f"Failed to process document: {str(e)}"}), 500
    finally:
        if os.path.exists(input_path):
            os.remove(input_path)

    original_name = os.path.splitext(file.filename)[0]
    return send_file(
        output_path,
        as_attachment=True,
        download_name=f"{original_name}_humanized.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    app.run(debug=False, host="0.0.0.0", port=int(os.environ.get("PORT", "5050")))
